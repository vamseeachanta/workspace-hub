"""Synthetic fixture generators for doc_intelligence tests.

All fixtures are generated programmatically — no binary files committed.
"""

import os
import tempfile
from pathlib import Path

import pytest


@pytest.fixture
def tmp_dir():
    """Provide a temporary directory that is cleaned up after the test."""
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


@pytest.fixture
def sample_pdf(tmp_dir):
    """Generate a minimal PDF with text, a table, and a figure reference."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=False)

    # Page 1 — heading + body text + figure reference
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Chapter 1: Introduction", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 12)
    pdf.multi_cell(0, 6, (
        "This document covers riser design fundamentals. "
        "See Figure 3.2 for the general arrangement."
    ))

    # Page 2 — a simple table
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "Section 2: Data", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 12)
    headers = ["Parameter", "Value", "Unit"]
    col_w = 50
    for h in headers:
        pdf.cell(col_w, 8, h, border=1)
    pdf.ln()
    for row in [("Diameter", "12", "inch"), ("Length", "500", "m")]:
        for val in row:
            pdf.cell(col_w, 8, val, border=1)
        pdf.ln()

    out = tmp_dir / "sample.pdf"
    pdf.output(str(out))
    return out


@pytest.fixture
def sample_docx(tmp_dir):
    """Generate a minimal DOCX with headings, body, and a table."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401

    doc = Document()
    doc.add_heading("Chapter 1: Overview", level=1)
    doc.add_paragraph("This section provides an overview of the design.")
    doc.add_heading("1.1 Scope", level=2)
    doc.add_paragraph("The scope covers subsea systems.")

    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "Riser"
    table.cell(1, 1).text = "Complete"
    table.cell(2, 0).text = "Mooring"
    table.cell(2, 1).text = "Pending"

    out = tmp_dir / "sample.docx"
    doc.save(str(out))
    return out


@pytest.fixture
def sample_xlsx(tmp_dir):
    """Generate a minimal XLSX with two sheets."""
    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Parameters"
    ws1.append(["Name", "Value", "Unit"])
    ws1.append(["Diameter", 12.0, "inch"])
    ws1.append(["Length", 500.0, "m"])

    ws2 = wb.create_sheet("Summary")
    ws2.append(["Metric", "Result"])
    ws2.append(["Total", 42])
    # Add an empty row to test skip-empty logic
    ws2.append([None, None])
    ws2.append(["Count", 7])

    out = tmp_dir / "sample.xlsx"
    wb.save(str(out))
    return out


@pytest.fixture
def corrupt_file(tmp_dir):
    """Create a file with garbage bytes for error-handling tests."""
    out = tmp_dir / "corrupt.pdf"
    out.write_bytes(b"\x00\x01\x02NOTAPDF\xff\xfe")
    return out
