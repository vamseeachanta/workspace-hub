#!/usr/bin/env python3
"""
Documentation to Context Converter
Converts various document formats to AI-friendly context files.
Leverages Claude Flow ecosystem for intelligent parsing.
"""

import os
import sys
import json
import mimetypes
import subprocess
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict
from datetime import datetime
import hashlib

# Document format detection
import magic  # python-magic for MIME type detection

# PDF parsing
try:
    import PyPDF2
    import pdfplumber
    from pdf2image import convert_from_path
    import pytesseract
    from PIL import Image
    HAS_PDF_SUPPORT = True
except ImportError:
    HAS_PDF_SUPPORT = False

# Word/DOCX parsing
try:
    from docx import Document as DocxDocument
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    HAS_DOCX_SUPPORT = True
except ImportError:
    HAS_DOCX_SUPPORT = False

# Excel parsing
try:
    import openpyxl
    from openpyxl.formula import Tokenizer
    HAS_EXCEL_SUPPORT = True
except ImportError:
    HAS_EXCEL_SUPPORT = False

# HTML parsing
try:
    from bs4 import BeautifulSoup
    import html2text
    HAS_HTML_SUPPORT = True
except ImportError:
    HAS_HTML_SUPPORT = False

# Markdown generation
import textwrap


@dataclass
class DocumentMetadata:
    """Metadata extracted from documents."""
    filename: str
    format: str
    mime_type: str
    size_bytes: int
    created_at: Optional[str] = None
    modified_at: Optional[str] = None
    author: Optional[str] = None
    title: Optional[str] = None
    page_count: Optional[int] = None
    sheet_count: Optional[int] = None
    word_count: Optional[int] = None
    checksum: Optional[str] = None
    extraction_timestamp: str = None

    def __post_init__(self):
        if self.extraction_timestamp is None:
            self.extraction_timestamp = datetime.now().isoformat()


@dataclass
class DocumentContent:
    """Parsed document content."""
    metadata: DocumentMetadata
    text: str
    structure: Dict[str, Any]
    tables: List[Dict[str, Any]]
    formulas: List[Dict[str, str]]
    images: List[Dict[str, str]]
    links: List[Dict[str, str]]
    raw_data: Optional[Dict[str, Any]] = None


class DocumentParser:
    """Base class for document parsers."""

    @staticmethod
    def can_handle(mime_type: str, extension: str) -> bool:
        """Check if this parser can handle the document type."""
        raise NotImplementedError

    def parse(self, file_path: Path) -> DocumentContent:
        """Parse document and return structured content."""
        raise NotImplementedError


class PDFParser(DocumentParser):
    """Parser for PDF documents (readable and scanned)."""

    @staticmethod
    def can_handle(mime_type: str, extension: str) -> bool:
        return mime_type == 'application/pdf' or extension == '.pdf'

    def parse(self, file_path: Path) -> DocumentContent:
        """Parse PDF with automatic OCR fallback for scanned documents."""
        if not HAS_PDF_SUPPORT:
            raise ImportError("PDF support not available. Install: pip install PyPDF2 pdfplumber pdf2image pytesseract pillow")

        metadata = self._extract_metadata(file_path)
        text_content = []
        tables = []
        images = []

        # Try text extraction first
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_content.append(f"### Page {page_num}\n\n{page_text}\n")

                    # Extract tables
                    page_tables = page.extract_tables()
                    if page_tables:
                        for table_num, table in enumerate(page_tables, 1):
                            tables.append({
                                'page': page_num,
                                'table_number': table_num,
                                'data': table,
                                'markdown': self._table_to_markdown(table)
                            })
        except Exception as e:
            print(f"Text extraction failed: {e}, falling back to OCR", file=sys.stderr)

        # If no text extracted, use OCR
        if not text_content or len(''.join(text_content).strip()) < 100:
            text_content = self._ocr_pdf(file_path)

        full_text = '\n\n'.join(text_content)

        return DocumentContent(
            metadata=metadata,
            text=full_text,
            structure={'pages': len(text_content)},
            tables=tables,
            formulas=[],
            images=images,
            links=[]
        )

    def _extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract PDF metadata."""
        stat = file_path.stat()

        metadata = DocumentMetadata(
            filename=file_path.name,
            format='PDF',
            mime_type='application/pdf',
            size_bytes=stat.st_size,
            created_at=datetime.fromtimestamp(stat.st_ctime).isoformat(),
            modified_at=datetime.fromtimestamp(stat.st_mtime).isoformat(),
            checksum=self._compute_checksum(file_path)
        )

        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                metadata.page_count = len(pdf_reader.pages)

                if pdf_reader.metadata:
                    metadata.author = pdf_reader.metadata.get('/Author')
                    metadata.title = pdf_reader.metadata.get('/Title')
        except Exception as e:
            print(f"Could not extract PDF metadata: {e}", file=sys.stderr)

        return metadata

    def _ocr_pdf(self, file_path: Path) -> List[str]:
        """Perform OCR on scanned PDF."""
        text_content = []
        try:
            images = convert_from_path(file_path)
            for page_num, image in enumerate(images, 1):
                text = pytesseract.image_to_string(image)
                if text.strip():
                    text_content.append(f"### Page {page_num} (OCR)\n\n{text}\n")
        except Exception as e:
            print(f"OCR failed: {e}", file=sys.stderr)
            text_content.append("### OCR Failed\n\nCould not extract text from scanned PDF.")

        return text_content

    @staticmethod
    def _table_to_markdown(table: List[List]) -> str:
        """Convert table data to markdown format."""
        if not table or not table[0]:
            return ""

        lines = []
        # Header
        lines.append("| " + " | ".join(str(cell or '') for cell in table[0]) + " |")
        lines.append("| " + " | ".join("---" for _ in table[0]) + " |")

        # Rows
        for row in table[1:]:
            lines.append("| " + " | ".join(str(cell or '') for cell in row) + " |")

        return "\n".join(lines)

    @staticmethod
    def _compute_checksum(file_path: Path) -> str:
        """Compute SHA-256 checksum of file."""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()


class WordParser(DocumentParser):
    """Parser for Word/DOCX documents."""

    @staticmethod
    def can_handle(mime_type: str, extension: str) -> bool:
        return (mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                or extension == '.docx')

    def parse(self, file_path: Path) -> DocumentContent:
        """Parse Word document preserving structure."""
        if not HAS_DOCX_SUPPORT:
            raise ImportError("DOCX support not available. Install: pip install python-docx")

        doc = DocxDocument(file_path)
        metadata = self._extract_metadata(file_path, doc)

        text_parts = []
        tables = []
        structure = {'sections': [], 'headings': []}

        # Parse document elements in order
        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                text = para.text.strip()
                if text:
                    # Check if heading
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        markdown_heading = '#' * int(level) if level.isdigit() else '##'
                        text_parts.append(f"{markdown_heading} {text}\n")
                        structure['headings'].append({'level': level, 'text': text})
                    else:
                        text_parts.append(f"{text}\n")

            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                table_data = self._extract_table(table)
                tables.append({
                    'table_number': len(tables) + 1,
                    'data': table_data,
                    'markdown': self._table_to_markdown(table_data)
                })
                text_parts.append(f"\n[Table {len(tables)}]\n")

        full_text = '\n'.join(text_parts)

        return DocumentContent(
            metadata=metadata,
            text=full_text,
            structure=structure,
            tables=tables,
            formulas=[],
            images=[],
            links=[]
        )

    def _extract_metadata(self, file_path: Path, doc: DocxDocument) -> DocumentMetadata:
        """Extract Word document metadata."""
        stat = file_path.stat()
        core_props = doc.core_properties

        # Count words
        word_count = sum(len(para.text.split()) for para in doc.paragraphs)

        return DocumentMetadata(
            filename=file_path.name,
            format='DOCX',
            mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            size_bytes=stat.st_size,
            created_at=core_props.created.isoformat() if core_props.created else None,
            modified_at=core_props.modified.isoformat() if core_props.modified else None,
            author=core_props.author,
            title=core_props.title,
            word_count=word_count,
            checksum=PDFParser._compute_checksum(file_path)
        )

    @staticmethod
    def _extract_table(table: Table) -> List[List[str]]:
        """Extract table data from Word table."""
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        return data

    @staticmethod
    def _table_to_markdown(table_data: List[List[str]]) -> str:
        """Convert table to markdown."""
        return PDFParser._table_to_markdown(table_data)


class ExcelParser(DocumentParser):
    """Parser for Excel spreadsheets."""

    @staticmethod
    def can_handle(mime_type: str, extension: str) -> bool:
        return (mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                or extension in ['.xlsx', '.xlsm'])

    def parse(self, file_path: Path) -> DocumentContent:
        """Parse Excel file with formulas and data."""
        if not HAS_EXCEL_SUPPORT:
            raise ImportError("Excel support not available. Install: pip install openpyxl")

        wb = openpyxl.load_workbook(file_path, data_only=False)
        metadata = self._extract_metadata(file_path, wb)

        text_parts = []
        tables = []
        formulas = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]

            # Skip chart sheets - they don't have data rows
            if not hasattr(sheet, 'iter_rows'):
                text_parts.append(f"## Sheet: {sheet_name} (Chart Sheet - Skipped)\n")
                continue

            text_parts.append(f"## Sheet: {sheet_name}\n")

            # Extract data and formulas
            sheet_data = []
            for row in sheet.iter_rows():
                row_data = []
                for cell in row:
                    value = cell.value
                    row_data.append(str(value) if value is not None else '')

                    # Check for formulas
                    if isinstance(value, str) and value.startswith('='):
                        formulas.append({
                            'sheet': sheet_name,
                            'cell': cell.coordinate,
                            'formula': value,
                            'value': str(cell.value)
                        })

                if any(row_data):  # Skip empty rows
                    sheet_data.append(row_data)

            if sheet_data:
                tables.append({
                    'sheet': sheet_name,
                    'table_number': len(tables) + 1,
                    'data': sheet_data,
                    'markdown': self._table_to_markdown(sheet_data)
                })

        # Add formula summary
        if formulas:
            text_parts.append(f"\n### Formulas Found ({len(formulas)} total)\n")
            for formula in formulas[:20]:  # Limit to first 20
                text_parts.append(f"- **{formula['cell']}** ({formula['sheet']}): `{formula['formula']}`")
            if len(formulas) > 20:
                text_parts.append(f"\n... and {len(formulas) - 20} more formulas\n")

        full_text = '\n'.join(text_parts)

        return DocumentContent(
            metadata=metadata,
            text=full_text,
            structure={'sheets': wb.sheetnames},
            tables=tables,
            formulas=formulas,
            images=[],
            links=[]
        )

    def _extract_metadata(self, file_path: Path, workbook) -> DocumentMetadata:
        """Extract Excel metadata."""
        stat = file_path.stat()
        props = workbook.properties

        return DocumentMetadata(
            filename=file_path.name,
            format='XLSX',
            mime_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            size_bytes=stat.st_size,
            created_at=props.created.isoformat() if props.created else None,
            modified_at=props.modified.isoformat() if props.modified else None,
            author=props.creator,
            title=props.title,
            sheet_count=len(workbook.sheetnames),
            checksum=PDFParser._compute_checksum(file_path)
        )

    @staticmethod
    def _table_to_markdown(table_data: List[List[str]]) -> str:
        """Convert table to markdown."""
        return PDFParser._table_to_markdown(table_data)


class HTMLParser(DocumentParser):
    """Parser for HTML documents."""

    @staticmethod
    def can_handle(mime_type: str, extension: str) -> bool:
        return mime_type == 'text/html' or extension in ['.html', '.htm']

    def parse(self, file_path: Path) -> DocumentContent:
        """Parse HTML document."""
        if not HAS_HTML_SUPPORT:
            raise ImportError("HTML support not available. Install: pip install beautifulsoup4 html2text")

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')
        metadata = self._extract_metadata(file_path, soup)

        # Convert to markdown
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = False
        h.ignore_emphasis = False
        markdown_text = h.handle(html_content)

        # Extract tables
        tables = []
        for idx, table in enumerate(soup.find_all('table'), 1):
            table_data = self._extract_html_table(table)
            tables.append({
                'table_number': idx,
                'data': table_data,
                'markdown': PDFParser._table_to_markdown(table_data)
            })

        # Extract links
        links = []
        for link in soup.find_all('a', href=True):
            links.append({
                'text': link.get_text(strip=True),
                'url': link['href']
            })

        return DocumentContent(
            metadata=metadata,
            text=markdown_text,
            structure={'title': soup.title.string if soup.title else None},
            tables=tables,
            formulas=[],
            images=[],
            links=links
        )

    def _extract_metadata(self, file_path: Path, soup: BeautifulSoup) -> DocumentMetadata:
        """Extract HTML metadata."""
        stat = file_path.stat()

        # Try to get title
        title = None
        if soup.title:
            title = soup.title.string

        # Try to get author from meta tags
        author = None
        author_meta = soup.find('meta', attrs={'name': 'author'})
        if author_meta:
            author = author_meta.get('content')

        # Count words in visible text
        text = soup.get_text()
        word_count = len(text.split())

        return DocumentMetadata(
            filename=file_path.name,
            format='HTML',
            mime_type='text/html',
            size_bytes=stat.st_size,
            created_at=datetime.fromtimestamp(stat.st_ctime).isoformat(),
            modified_at=datetime.fromtimestamp(stat.st_mtime).isoformat(),
            author=author,
            title=title,
            word_count=word_count,
            checksum=PDFParser._compute_checksum(file_path)
        )

    @staticmethod
    def _extract_html_table(table_element) -> List[List[str]]:
        """Extract table data from HTML table element."""
        data = []
        for row in table_element.find_all('tr'):
            row_data = []
            for cell in row.find_all(['td', 'th']):
                row_data.append(cell.get_text(strip=True))
            if row_data:
                data.append(row_data)
        return data


class DocumentToContextConverter:
    """Main converter class that orchestrates document parsing."""

    def __init__(self):
        self.parsers = [
            PDFParser(),
            WordParser(),
            ExcelParser(),
            HTMLParser()
        ]

    def detect_format(self, file_path: Path) -> Tuple[str, str]:
        """Detect document format using magic and extension."""
        # Get MIME type
        mime = magic.Magic(mime=True)
        mime_type = mime.from_file(str(file_path))

        # Get extension
        extension = file_path.suffix.lower()

        return mime_type, extension

    def select_parser(self, file_path: Path) -> Optional[DocumentParser]:
        """Select appropriate parser for document."""
        mime_type, extension = self.detect_format(file_path)

        for parser in self.parsers:
            if parser.can_handle(mime_type, extension):
                return parser

        return None

    def convert(self, input_path: str, output_path: Optional[str] = None,
                output_format: str = 'markdown') -> DocumentContent:
        """Convert document to context file."""
        file_path = Path(input_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {input_path}")

        # Select parser
        parser = self.select_parser(file_path)
        if not parser:
            mime_type, extension = self.detect_format(file_path)
            raise ValueError(f"No parser available for {mime_type} ({extension})")

        print(f"Using {parser.__class__.__name__} for {file_path.name}", file=sys.stderr)

        # Parse document
        content = parser.parse(file_path)

        # Generate output
        if output_path:
            output_file = Path(output_path)
        else:
            output_file = file_path.with_suffix('.context.md')

        if output_format == 'markdown':
            self._write_markdown(content, output_file)
        elif output_format == 'json':
            self._write_json(content, output_file.with_suffix('.json'))
        else:
            raise ValueError(f"Unsupported output format: {output_format}")

        print(f"Context file created: {output_file}", file=sys.stderr)
        return content

    def _write_markdown(self, content: DocumentContent, output_path: Path):
        """Write content as markdown context file."""
        with open(output_path, 'w', encoding='utf-8') as f:
            # Write header
            f.write("---\n")
            f.write("# Document Context\n")
            f.write(f"Generated: {content.metadata.extraction_timestamp}\n")
            f.write("---\n\n")

            # Write metadata
            f.write("## Document Metadata\n\n")
            metadata_dict = asdict(content.metadata)
            for key, value in metadata_dict.items():
                if value is not None:
                    f.write(f"- **{key.replace('_', ' ').title()}**: {value}\n")
            f.write("\n")

            # Write main content
            f.write("## Document Content\n\n")
            f.write(content.text)
            f.write("\n\n")

            # Write tables
            if content.tables:
                f.write("## Tables\n\n")
                for table in content.tables:
                    f.write(f"### Table {table.get('table_number', '?')}")
                    if 'sheet' in table:
                        f.write(f" (Sheet: {table['sheet']})")
                    if 'page' in table:
                        f.write(f" (Page: {table['page']})")
                    f.write("\n\n")
                    f.write(table['markdown'])
                    f.write("\n\n")

            # Write formulas
            if content.formulas:
                f.write("## Formulas\n\n")
                for formula in content.formulas:
                    f.write(f"- **{formula.get('cell', '?')}** ")
                    if 'sheet' in formula:
                        f.write(f"({formula['sheet']})")
                    f.write(f": `{formula.get('formula', '?')}`\n")
                f.write("\n")

            # Write links
            if content.links:
                f.write("## Links\n\n")
                for link in content.links[:50]:  # Limit to 50 links
                    f.write(f"- [{link.get('text', 'Link')}]({link.get('url', '#')})\n")
                if len(content.links) > 50:
                    f.write(f"\n... and {len(content.links) - 50} more links\n")
                f.write("\n")

    def _write_json(self, content: DocumentContent, output_path: Path):
        """Write content as JSON."""
        output_data = {
            'metadata': asdict(content.metadata),
            'text': content.text,
            'structure': content.structure,
            'tables': content.tables,
            'formulas': content.formulas,
            'images': content.images,
            'links': content.links
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False)


def main():
    """CLI entry point."""
    import argparse

    parser = argparse.ArgumentParser(
        description='Convert documents to AI-friendly context files',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent('''
        Supported formats:
          - PDF (readable and scanned with OCR)
          - Word/DOCX
          - Excel/XLSX (with formula extraction)
          - HTML

        Examples:
          %(prog)s document.pdf
          %(prog)s report.docx -o context.md
          %(prog)s data.xlsx -f json
          %(prog)s *.pdf -b
        ''')
    )

    parser.add_argument('input', nargs='+', help='Input document(s) to convert')
    parser.add_argument('-o', '--output', help='Output file path (default: input.context.md)')
    parser.add_argument('-f', '--format', choices=['markdown', 'json'],
                       default='markdown', help='Output format (default: markdown)')
    parser.add_argument('-b', '--batch', action='store_true',
                       help='Batch mode: process multiple files')

    args = parser.parse_args()

    converter = DocumentToContextConverter()

    # Expand glob patterns
    import glob
    input_files = []
    for pattern in args.input:
        matches = glob.glob(pattern)
        if matches:
            input_files.extend(matches)
        else:
            input_files.append(pattern)

    # Process files
    for input_file in input_files:
        try:
            output_file = args.output if not args.batch else None
            converter.convert(input_file, output_file, args.format)
        except Exception as e:
            print(f"Error processing {input_file}: {e}", file=sys.stderr)
            if not args.batch:
                sys.exit(1)


if __name__ == '__main__':
    main()
