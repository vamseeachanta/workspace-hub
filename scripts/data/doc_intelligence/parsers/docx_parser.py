"""DOCX parser using python-docx for heading/body/table extraction."""

import hashlib
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import List

from scripts.data.doc_intelligence.parsers.base import BaseParser
from scripts.data.doc_intelligence.schema import (
    DocumentManifest,
    DocumentMetadata,
    ExtractedSection,
    ExtractedTable,
    SourceLocation,
)


def _compute_checksum(filepath: str) -> str:
    sha = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            sha.update(chunk)
    return sha.hexdigest()


def _heading_level(style_name: str) -> int:
    """Extract heading level from style name like 'Heading 1' → 1."""
    m = re.match(r"Heading\s+(\d+)", style_name or "", re.IGNORECASE)
    return int(m.group(1)) if m else 0


class DocxParser(BaseParser):
    """Extract headings, body paragraphs, and tables from DOCX files."""

    def can_handle(self, filepath: str) -> bool:
        return Path(filepath).suffix.lower() == ".docx"

    def parse(self, filepath: str, domain: str) -> DocumentManifest:
        p = Path(filepath)
        meta = DocumentMetadata(
            filename=p.name,
            format="docx",
            size_bytes=p.stat().st_size,
            checksum=_compute_checksum(filepath),
            extraction_timestamp=datetime.now(timezone.utc).isoformat(),
        )
        sections: List[ExtractedSection] = []
        tables: List[ExtractedTable] = []
        errors: List[str] = []

        try:
            from docx import Document

            doc = Document(str(p))
            current_heading: str | None = None

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style = para.style.name if para.style else ""
                level = _heading_level(style)
                if level > 0:
                    current_heading = text
                sections.append(
                    ExtractedSection(
                        heading=text if level > 0 else current_heading,
                        level=level,
                        text=text,
                        source=SourceLocation(
                            document=p.name, section=current_heading
                        ),
                    )
                )

            for tbl in doc.tables:
                rows_data = []
                for row in tbl.rows:
                    rows_data.append([cell.text.strip() for cell in row.cells])
                if not rows_data:
                    continue
                header = rows_data[0]
                data_rows = rows_data[1:] if len(rows_data) > 1 else []
                tables.append(
                    ExtractedTable(
                        title=None,
                        columns=header,
                        rows=data_rows,
                        source=SourceLocation(document=p.name),
                    )
                )
        except Exception as exc:
            errors.append(f"DOCX extraction failed: {exc}")

        return DocumentManifest(
            version="1.0.0",
            tool="extract-document/1.0.0",
            domain=domain,
            metadata=meta,
            sections=sections,
            tables=tables,
            figure_refs=[],
            extraction_stats={
                "sections": len(sections),
                "tables": len(tables),
                "figure_refs": 0,
            },
            errors=errors,
        )
